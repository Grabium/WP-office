from docx import Document
from datetime import datetime
import sys
import os

"""
Este script captura um documento word.docx pronto e troca textos em tabela e parágrafos
"""

#sys.argv é uma lista que recebe todos os argumentos
#parametros com valores  "<arg>" atrapalham o recebimento em python
#não é possível receber um array. Somente strings e formatá-las como array


#recebendo os argumentos e transformando-os num dicionário
referencias = {}
indice = ''
for key, item in enumerate(sys.argv):#enumerate salva o índice em key.
  if key == 0: #nome do arquivo main.py
    continue
  if ((key % 2) != 0):
    indice = f"<{item}>"
    referencias[f"<{item}>"] = ''
  else:
    referencias[indice] = item
    
#print (referencias)
#sys.exit('teste de recebimento de argumentos vindo do php: ok')


referencias['<dia>'] = str(datetime.now().day)
referencias['<mes>'] = str(datetime.now().month)
referencias['<ano>'] = str(datetime.now().year)


#print (referencias)
#sys.exit('teste de recebimento de argumentos vindo do php: ok')

# O dicionário fica assim:
#referencias = {
#  '<port_num>' : port_num,
#  '<l_is>': l_is,
#  '<s>':s,
#  '<contrato_num>':contrato_num,
#  '<empresa_nome>':empresa_nome,
#  '<empresa_cnpj>':empresa_cnpj,
#  '<posto>':posto,
#  '<CONTRATO_OBJETO>':contrato_ojbeto,
#  '<titular_nome>':titular_nome,
#  '<titular_matricula>':titular_matricula,
#  '<substituto_nome>':substituto_nome,
#  '<substituto_matricula>':substituto_matricula,
#  '<assinante_nome>':assinante_nome,
#  '<assinante_cargo>':assinante_cargo,
#  '<dia>'  : dia,
#  '<mes>'  : mes,
#  '<ano>'  : ano
#}


#buscando o documento
document = Document('portaria_model_quick_docs.docx')


#descompactando os elementos tabelas até a camada de texto.
tables = document.tables #acho que são todas as tabelas do documento.
last_cell = ''
for table in tables: 
  for cell in table._cells: #são as células de uma tabela só.
    if last_cell is cell : #células que ocupam mais de uma unidade-cubo se repetem
      continue
    last_cell = cell
    #print(f'celula :{cell} - texto: {cell.text}.')
    for key, valor in referencias.items():#items() faz a relação: key, valor funcionar.
      #print(f'itera: {key}: {valor}')
      cell.text = (cell.text).replace(key, valor)# string.replace(novo, velho)
    #print(f'formatou :{cell.text}')

#for table in tables: 
  #for cell in table._cells: #são as células de uma tabela só.
    #print(f'Resultado :{cell.text}:')
    #sys.exit('fim')
#sys.exit('teste de substituição de tabela: ok')



#descompactando os elementos parágrafos até a camada de texo.
#print(document.paragraphs)#exibe uma array com os ponteiros dos paragrafos.
for paragrafo in document.paragraphs:
  for key, valor in referencias.items():
    paragrafo.text = paragrafo.text.replace(key, valor)
  
#for item in document.paragraphs:
  #print(item.text)

#contando a quantidade de documentos já salvos e adicionando um sufixo numérico de contagem ao nome do documento.
n = (len(os.listdir('novos_documentos')) + 1)

document.save(f'novos_documentos/novo({n}).docx')
#document.save(f"novos_documentos/{sys.argv[1]}.docx")